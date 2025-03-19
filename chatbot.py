import os
import sqlite3
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from telegram import Update
from telegram.ext import Application, CommandHandler, MessageHandler, filters, CallbackContext

def init_db():
    conn = sqlite3.connect('C:\\ChatMetodicBot\\database.db')  # Указываем путь к базе данных
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS topics (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT UNIQUE,
            content_path TEXT
        )
    ''')
    c.execute('''
        CREATE TABLE IF NOT EXISTS subtopics (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            topic_id INTEGER,
            name TEXT,
            content TEXT,
            order_num INTEGER,
            FOREIGN KEY (topic_id) REFERENCES topics(id)
        )
    ''')
    conn.commit()
    conn.close()

def extract_text_from_word(file_path):
    doc = Document(file_path)
    subtopics = []
    order_num = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("*b*") and "**" in text and "*e*" in text:
            parts = text.split("**")
            if len(parts) == 2:
                order_num += 1
                name = parts[0].replace("*b*", "").strip()
                content = parts[1].replace("*e*", "").strip()
                subtopic = {"name": name, "content": content, "order_num": order_num}
                subtopics.append(subtopic)

    return subtopics

def add_topic_with_subtopics(topic_name, subtopics, file_path):
    conn = sqlite3.connect('C:\\ChatMetodicBot\\database.db')
    c = conn.cursor()
    c.execute("INSERT OR IGNORE INTO topics (name, content_path) VALUES (?, ?)", (topic_name, file_path))
    c.execute("SELECT id FROM topics WHERE name=?", (topic_name,))
    topic_id = c.fetchone()[0]

    for subtopic in subtopics:
        c.execute("INSERT INTO subtopics (topic_id, name, content, order_num) VALUES (?, ?, ?, ?)",
                  (topic_id, subtopic["name"], subtopic["content"], subtopic["order_num"]))

    conn.commit()
    conn.close()

def load_files_from_directory(directory):
    conn = sqlite3.connect('C:\\ChatMetodicBot\\database.db')
    c = conn.cursor()
    c.execute("DELETE FROM topics")
    c.execute("DELETE FROM subtopics")
    conn.commit()

    for filename in os.listdir(directory):
        if filename.endswith(".docx") and not filename.startswith("~$"):
            file_path = os.path.join(directory, filename)
            subtopics = extract_text_from_word(file_path)
            topic_name = os.path.splitext(filename)[0]  # Используем имя файла как имя темы
            add_topic_with_subtopics(topic_name, subtopics, file_path)
            print(f"Файл {filename} добавлен в базу данных.")

    conn.close()

def clean_paragraph(paragraph):
    text = paragraph.text
    if "*b*" in text and "**" in text and "*e*" in text:
        text = text.replace("*b*", "").replace("**", "").replace("*e*", "").strip()
    return text

async def start(update: Update, context: CallbackContext):
    print("Команда /start была получена")
    conn = sqlite3.connect('C:\\ChatMetodicBot\\database.db')
    c = conn.cursor()
    c.execute("SELECT id, name FROM topics ORDER BY name")
    topics = c.fetchall()

    message = "Привет! Выбери темы, которые тебя интересуют:\n"
    for idx, topic in enumerate(topics, start=1):
        topic_id, topic_name = topic
        message += f"{idx}: {topic_name}"
        c.execute("SELECT id, name FROM subtopics WHERE topic_id=? ORDER BY order_num", (topic_id,))
        subtopics = c.fetchall()
        for i, subtopic in enumerate(subtopics, start=1):
            sub_id, sub_name = subtopic
            message += f" {idx}.{i}: {sub_name}"
        message += "\n"

    await update.message.reply_text(message)
    conn.close()

async def handle_message(update: Update, context: CallbackContext):
    print("Сообщение от пользователя получено")
    selected_items = update.message.text.split('+')

    conn = sqlite3.connect('C:\\ChatMetodicBot\\database.db')
    c = conn.cursor()

    doc = Document()

    for item in selected_items:
        if '.' in item:
            topic_idx, sub_idx = map(int, item.split('.'))
            c.execute("SELECT id FROM topics ORDER BY name")
            topics = c.fetchall()
            topic_id = topics[topic_idx - 1][0]
            c.execute("SELECT name, content FROM subtopics WHERE topic_id=? ORDER BY order_num", (topic_id,))
            subtopics = c.fetchall()
            subtopic = subtopics[sub_idx - 1]
            subtopic_name = subtopic[0]
            subtopic_content = subtopic[1]
            doc.add_heading(subtopic_name, level=2)
            doc.add_paragraph(subtopic_content)
        else:
            topic_idx = int(item)
            c.execute("SELECT id FROM topics ORDER BY name")
            topics = c.fetchall()
            topic_id = topics[topic_idx - 1][0]
            c.execute("SELECT name, content_path FROM topics WHERE id=?", (topic_id,))
            topic = c.fetchone()
            topic_name = topic[0]
            topic_content_path = topic[1]
            doc.add_heading(topic_name, level=1)
            original_doc = Document(topic_content_path)
            for para in original_doc.paragraphs:
                cleaned_text = clean_paragraph(para)
                doc.add_paragraph(cleaned_text)

    output_dir = 'C:\\ChatMetodicBot\\ComplexFile'
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, 'new_file.docx')
    doc.save(output_path)
        
    await update.message.reply_text("Вот твой файл.")
    with open(output_path, 'rb') as f:
        await update.message.reply_document(f)
    
    conn.close()

if __name__ == "__main__":
    init_db()
    load_files_from_directory("C:\\ChatMetodicBot")
    application = Application.builder().token("7644850099:AAFT4Od8732MvEHv7A-290VNssU21UTeZvc").build()
    application.add_handler(CommandHandler("start", start))
    application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message))
    print("Бот запущен и готов к приему команд")
    application.run_polling()
